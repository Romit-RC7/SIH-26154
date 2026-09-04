"""
Pytest configuration and shared fixtures for SIH-26154 testing suite.
Uses SQLite in-memory database with SQLAlchemy async engine to test without requiring external DB.
"""

import io
from pathlib import Path
from typing import AsyncGenerator
import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker, AsyncSession
import fitz
import docx

from backend.app.core.config import settings
from backend.app.database.base import Base
from backend.app.database.session import get_db
from backend.app.main import app

# Test In-Memory SQLite Async Database
TEST_DB_URL = "sqlite+aiosqlite:///:memory:"

test_engine = create_async_engine(
    TEST_DB_URL,
    echo=False,
    connect_args={"check_same_thread": False},
)

TestingSessionLocal = async_sessionmaker(
    bind=test_engine,
    class_=AsyncSession,
    expire_on_commit=False,
    autocommit=False,
    autoflush=False,
)


from backend.app.services.pipeline_service import pipeline_service


@pytest_asyncio.fixture(scope="function")
async def db_session() -> AsyncGenerator[AsyncSession, None]:
    """Provides a clean in-memory database session for each test."""
    old_factory = pipeline_service.session_factory
    pipeline_service.session_factory = TestingSessionLocal

    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    async with TestingSessionLocal() as session:
        yield session

    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)

    pipeline_service.session_factory = old_factory


@pytest_asyncio.fixture(scope="function")
async def client(db_session: AsyncSession) -> AsyncGenerator[AsyncClient, None]:
    """Configures an asynchronous test client with overridden database dependency."""
    async def override_get_db():
        yield db_session

    app.dependency_overrides[get_db] = override_get_db

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://testserver") as ac:
        yield ac

    app.dependency_overrides.clear()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Generates a valid test PDF document in memory using PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)

    # Insert Title
    page.insert_text((50, 50), "AI Content Transformation Architecture", fontsize=18)

    # Insert Paragraph
    page.insert_text((50, 100), "This is an automated test document verifying Phase 1 of SIH-26154.", fontsize=11)
    page.insert_text((50, 120), "It covers document parsing, layout recognition, and semantic contract generation.", fontsize=11)

    # Insert a Table block
    page.insert_text((50, 160), "Table 1: Benchmark Results", fontsize=12)
    page.insert_text((50, 185), "Metric | Baseline | Proposed", fontsize=10)
    page.insert_text((50, 205), "Accuracy | 88.4% | 96.2%", fontsize=10)
    page.insert_text((50, 225), "Latency | 450ms | 120ms", fontsize=10)

    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Generates a valid test DOCX document in memory using python-docx."""
    doc = docx.Document()
    doc.add_heading("Semantic Document Pipeline Test", level=1)
    doc.add_paragraph("Testing DOCX ingestion into the Semantic Document JSON structure.")

    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Feature"
    hdr_cells[1].text = "Status"

    row1 = table.rows[1].cells
    row1[0].text = "PP-StructureV3"
    row1[1].text = "Integrated"

    row2 = table.rows[2].cells
    row2[0].text = "Semantic Contract"
    row2[1].text = "Verified"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
