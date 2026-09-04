"""
Extended Unit Tests for PDF and DOCX parsers covering edge cases:
- Multi-page PDF
- Empty/blank PDF
- Corrupted PDF
- Large PDF
- Complex DOCX
- Corrupted DOCX
"""

from pathlib import Path
import pytest
import pymupdf as fitz
import docx

from backend.app.processors.pdf_parser import pdf_parser
from backend.app.processors.docx_parser import docx_parser


def test_multipage_pdf_parsing(tmp_path: Path):
    """Verify parsing a multi-page document across all pages."""
    doc = fitz.open()
    for i in range(1, 6):
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 50), f"Chapter {i}: Advanced Neural Architecture", fontsize=16)
        page.insert_text((50, 80), f"Detailed analysis and benchmarks for chapter {i}.", fontsize=10)

    pdf_file = tmp_path / "multipage.pdf"
    doc.save(str(pdf_file))
    doc.close()

    pages, raw_elements, metadata = pdf_parser.parse(pdf_file)

    assert len(pages) == 5
    assert metadata["page_count"] == 5
    assert len(raw_elements) >= 5
    for idx, p in enumerate(pages, 1):
        assert f"Chapter {idx}" in p.raw_text


def test_empty_blank_pdf(tmp_path: Path):
    """Verify parser handles a PDF with blank pages gracefully."""
    doc = fitz.open()
    doc.new_page(width=500, height=700)  # Blank page
    pdf_file = tmp_path / "blank.pdf"
    doc.save(str(pdf_file))
    doc.close()

    pages, raw_elements, metadata = pdf_parser.parse(pdf_file)
    assert len(pages) == 1
    assert metadata["page_count"] == 1
    assert len(raw_elements) == 0  # No text blocks in a blank page


def test_corrupted_pdf_handling(tmp_path: Path):
    """Verify error handling when encountering corrupted PDF file bytes."""
    corrupt_file = tmp_path / "corrupted.pdf"
    corrupt_file.write_bytes(b"NOT A VALID PDF HEADER %%% garbage content 12345")

    with pytest.raises(Exception):
        pdf_parser.parse(corrupt_file)


def test_nonexistent_pdf_handling(tmp_path: Path):
    """Verify FileNotFoundError when target PDF does not exist."""
    missing_file = tmp_path / "does_not_exist.pdf"
    with pytest.raises(FileNotFoundError):
        pdf_parser.parse(missing_file)


def test_large_pdf_parsing(tmp_path: Path):
    """Verify parser efficiently handles a larger document (15 pages with multiple paragraphs)."""
    doc = fitz.open()
    for i in range(1, 16):
        page = doc.new_page()
        page.insert_text((40, 40), f"Section {i} - Enterprise Data Flow", fontsize=14)
        for line_idx in range(5):
            page.insert_text((40, 70 + line_idx * 20), f"Paragraph {line_idx + 1}: Data validation and schema check.", fontsize=9)

    large_file = tmp_path / "large_doc.pdf"
    doc.save(str(large_file))
    doc.close()

    pages, raw_elements, metadata = pdf_parser.parse(large_file)
    assert len(pages) == 15
    assert metadata["page_count"] == 15
    assert len(raw_elements) >= 15


def test_complex_docx_parsing(tmp_path: Path):
    """Verify parsing DOCX with mixed headers, paragraphs, and asymmetric tables."""
    doc = docx.Document()
    doc.add_heading("Deep Learning System Overview", level=1)
    doc.add_paragraph("First paragraph describing multimodal capabilities.")
    doc.add_heading("Sub-section 2.1", level=2)
    doc.add_paragraph("Second paragraph with technical specifications.")

    # Add complex 4x3 table
    table = doc.add_table(rows=4, cols=3)
    headers = ["Model", "Parameters", "Accuracy"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h

    data = [
        ["Qwen2.5-VL", "3B", "94.2%"],
        ["PP-StructureV3", "Lightweight", "96.5%"],
        ["BGE-M3", "560M", "92.8%"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    docx_path = tmp_path / "complex.docx"
    doc.save(str(docx_path))

    pages, raw_elements, meta = docx_parser.parse(docx_path)
    assert len(raw_elements) >= 4
    table_elems = [e for e in raw_elements if e.type == "table"]
    assert len(table_elems) == 1
    assert "Qwen2.5-VL" in table_elems[0].markdown
    assert "<th>Parameters</th>" in table_elems[0].html
    assert table_elems[0].table_data["rows"][1][0] == "Qwen2.5-VL"


def test_corrupted_docx_handling(tmp_path: Path):
    """Verify error handling on corrupted DOCX binary."""
    corrupt_docx = tmp_path / "broken.docx"
    corrupt_docx.write_bytes(b"PK\x03\x04 corrupt zip archive data")

    with pytest.raises(Exception):
        docx_parser.parse(corrupt_docx)


def test_nonexistent_docx_handling(tmp_path: Path):
    """Verify FileNotFoundError on missing DOCX."""
    missing_docx = tmp_path / "missing.docx"
    with pytest.raises(FileNotFoundError):
        docx_parser.parse(missing_docx)
